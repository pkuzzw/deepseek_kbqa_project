import os
from openai import OpenAI
from docx import Document

# 设置API密钥和基础URL
api_key = "sk-ooscmypgomjmlrzejjhcidzhzegvmirvvweonaoacfmpcrrc"
base_url = "https://api.siliconflow.cn/v1"

# 初始化OpenAI客户端
client = OpenAI(api_key=api_key, base_url=base_url)

# 定义文件夹路径
folder_path = r"D:\AILearn"

# 定义模型的上下文长度限制
MAX_CONTEXT_LENGTH = 32000  # 32K characters

# 遍历文件夹中的所有docx文件
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        # 读取文档
        doc_path = os.path.join(folder_path, filename)
        print(f"正在读取文档: {doc_path}")
        doc = Document(doc_path)

        # 提取文档中的所有段落
        paragraphs = [p.text for p in doc.paragraphs]

        # 将段落内容拼接成一个字符串
        content = "\n".join(paragraphs)

        # 分块处理
        chunks = []
        current_chunk = ""
        current_length = 0

        for paragraph in paragraphs:
            if current_length + len(paragraph) + 1 > MAX_CONTEXT_LENGTH:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph
                current_length = len(paragraph)
            else:
                if current_chunk:
                    current_chunk += "\n" + paragraph
                else:
                    current_chunk = paragraph
                current_length += len(paragraph) + 1

        if current_chunk:
            chunks.append(current_chunk.strip())

        # 翻译每个块
        translated_chunks = []

        for i, chunk in enumerate(chunks):
            print(f"正在翻译文档: {filename} 的第 {i + 1}/{len(chunks)} 块")
            response = client.chat.completions.create(
                model='Qwen/Qwen2.5-7B-Instruct',
                messages=[
                    {'role': 'user', 'content': f"将下面的英文内容翻译成中文\n{chunk}"}
                    # "model": "Qwen/Qwen2.5-7B-Instruct",
                    # "messages": [
                    #     {"role": "system", "content": "你是一个知识问答助手，根据提供的背景信息回答问题。"},
                    #     {"role": "user", "content": f"背景信息：{context}"},
                    #     {"role": "user", "content": question}
                    # ],
                    # "temperature": 0.7,
                    # "max_tokens": 1024          
                ],
                stream=False
            )
            translated_chunks.append(response.choices[0].message.content)

        # 拼接翻译后的内容
        translated_content = "\n".join(translated_chunks)

        # 创建新的文档并写入翻译后的内容
        new_doc = Document()
        for line in translated_content.split("\n"):
            new_doc.add_paragraph(line)

        # 保存翻译后的文档
        new_filename = filename.replace(".docx", "_中文版.docx")
        new_doc_path = os.path.join(folder_path, new_filename)
        new_doc.save(new_doc_path)

        print(f"翻译完成，保存为: {new_doc_path}")

print("所有文档翻译完成！")
